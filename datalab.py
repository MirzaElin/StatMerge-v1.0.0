from PySide6 import QtWidgets, QtCore
try:
    from matplotlib.figure import Figure
    from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
except Exception:
    Figure=None; FigureCanvas=None
try:
    import pandas as pd
except Exception:
    pd=None
try:
    from sklearn.datasets import make_classification
except Exception:
    make_classification=None
try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document=None; Inches=None
class DataLab(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__(); self.setWindowTitle('Data Lab'); self.resize(1200,760); self.df=None; self._build()
    def _build(self):
        c=QtWidgets.QWidget(self); self.setCentralWidget(c); v=QtWidgets.QVBoxLayout(c)
        tb=QtWidgets.QToolBar('Tools'); self.addToolBar(tb)
        def add(text,slot): a=QtWidgets.QAction(text,self); a.triggered.connect(slot); tb.addAction(a)
        add('Load CSV',self.on_open); add('Save Clean CSV',self.on_save); tb.addSeparator(); add('Preview Head',self.on_head); add('Describe',self.on_desc); add('Correlation',self.on_corr); tb.addSeparator(); add('Plot',self.on_plot); tb.addSeparator(); add('Load Sample',self.on_sample); add('Export Report (DOCX)',self.on_report)
        sp=QtWidgets.QSplitter(self); sp.setOrientation(QtCore.Qt.Horizontal); v.addWidget(sp,1)
        self.txt=QtWidgets.QPlainTextEdit(self); self.txt.setReadOnly(True); sp.addWidget(self.txt)
        right=QtWidgets.QWidget(self); rv=QtWidgets.QVBoxLayout(right); sp.addWidget(right)
        self.table=QtWidgets.QTableWidget(self); rv.addWidget(self.table,3)
        if Figure and FigureCanvas:
            self.fig=Figure(figsize=(5,3)); self.canvas=FigureCanvas(self.fig); rv.addWidget(self.canvas,2)
        else:
            self.fig=None; self.canvas=None; lab=QtWidgets.QLabel('Matplotlib Qt backend not available — plots disabled',self); lab.setAlignment(QtCore.Qt.AlignCenter); rv.addWidget(lab,1)
    def _need(self):
        if pd is None: QtWidgets.QMessageBox.warning(self,'Pandas missing','Install pandas to use Data Lab.'); return False
        if self.df is None or getattr(self.df,'empty',True): QtWidgets.QMessageBox.information(self,'No data','Load a CSV or click Load Sample.'); return False
        return True
    def on_open(self):
        if pd is None: return
        p,_=QtWidgets.QFileDialog.getOpenFileName(self,'Open CSV','','CSV files (*.csv);;All files (*.*)')
        if not p: return
        try:
            self.df=pd.read_csv(p); self._fill(); self.txt.setPlainText('Loaded: '+p+'\nRows: '+str(len(self.df))+', Columns: '+str(len(self.df.columns)))
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,'Error',str(e))
    def on_save(self):
        if not self._need(): return
        p,_=QtWidgets.QFileDialog.getSaveFileName(self,'Save CSV','clean.csv','CSV files (*.csv)')
        if not p: return
        try:
            self.df.to_csv(p,index=False); QtWidgets.QMessageBox.information(self,'Saved','Saved to '+p)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,'Error',str(e))
    def _fill(self,max_rows=500):
        if self.df is None or pd is None: return
        dfv=self.df.head(max_rows); self.table.clear(); self.table.setRowCount(len(dfv)); self.table.setColumnCount(len(dfv.columns))
        self.table.setHorizontalHeaderLabels([str(c) for c in dfv.columns])
        for i in range(len(dfv)):
            for j in range(len(dfv.columns)):
                self.table.setItem(i,j,QtWidgets.QTableWidgetItem(str(dfv.iat[i,j])))
        self.table.resizeColumnsToContents()
    def on_head(self):
        if not self._need(): return
        self.txt.setPlainText(str(self.df.head(20)))
    def on_desc(self):
        if not self._need(): return
        try:
            desc=self.df.describe(include='all').transpose(); self.txt.setPlainText(str(desc))
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,'Error',str(e))
    def on_corr(self):
        if not self._need(): return
        try:
            num=self.df.select_dtypes(include='number')
            try: corr=num.corr(numeric_only=True)
            except TypeError: corr=num.corr()
            self.txt.setPlainText(str(corr))
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,'Error',str(e))
    def on_plot(self):
        if not self._need(): return
        if self.fig is None or self.canvas is None:
            QtWidgets.QMessageBox.information(self,'Plots disabled','Matplotlib Qt backend not available.'); return
        cols=list(self.df.columns)
        if not cols:
            QtWidgets.QMessageBox.information(self,'No columns','Dataset has no columns.'); return
        x,ok=QtWidgets.QInputDialog.getItem(self,'Plot X','Select X column:',cols,0,False)
        if not ok: return
        y,ok=QtWidgets.QInputDialog.getItem(self,'Plot Y (or same for hist)','Select Y column:',cols,0,False)
        if not ok: return
        self.fig.clear(); ax=self.fig.add_subplot(111)
        try:
            import pandas as _pd
            if x==y:
                ax.hist(_pd.to_numeric(self.df[x],errors='coerce').dropna()); ax.set_title('Histogram of '+x)
            else:
                ax.scatter(_pd.to_numeric(self.df[x],errors='coerce'),_pd.to_numeric(self.df[y],errors='coerce')); ax.set_xlabel(x); ax.set_ylabel(y); ax.set_title('Scatter: '+x+' vs '+y)
            self.canvas.draw()
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,'Plot error',str(e))
    def on_sample(self):
        if pd is None or make_classification is None:
            QtWidgets.QMessageBox.warning(self,'Missing libs','Install pandas and scikit-learn for sample data.'); return
        X,y=make_classification(n_samples=400,n_features=6,n_informative=4,random_state=42)
        import pandas as _pd
        cols=['f'+str(i) for i in range(len(X[0]))]; self.df=_pd.DataFrame(X,columns=cols); self.df['target']=y
        self._fill(); self.txt.setPlainText('Loaded a synthetic classification dataset (400x7).')
    def on_report(self):
        if not self._need(): return
        if Document is None:
            QtWidgets.QMessageBox.warning(self,'python-docx missing','Install python-docx to export reports.'); return
        import tempfile, datetime, pandas as _pd
        try:
            desc=self.df.describe(include='all').transpose()
            try: corr=self.df.select_dtypes(include='number').corr(numeric_only=True)
            except TypeError: corr=self.df.select_dtypes(include='number').corr()
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,'Error',str(e)); return
        now=datetime.datetime.now().strftime('%Y-%m-%d_%H%M%S')
        out_doc=os.path.join(os.path.expanduser('~'),'StatMerge_Report_'+now+'.docx')
        doc=Document(); doc.add_heading('StatMerge Data Lab Report',0); doc.add_paragraph('Generated: '+now); doc.add_paragraph('Rows: '+str(len(self.df))+', Columns: '+str(len(self.df.columns)))
        doc.add_heading('Summary statistics',1); doc.add_paragraph(str(desc))
        doc.add_heading('Correlation (numeric)',1); doc.add_paragraph(str(corr))
        if self.fig is not None and self.canvas is not None:
            try:
                png_path=os.path.join(tempfile.gettempdir(),'statmerge_plot_'+now+'.png')
                if not self.fig.axes:
                    num_cols=list(self.df.select_dtypes(include='number').columns)
                    if num_cols:
                        self.fig.clear(); ax=self.fig.add_subplot(111); ax.hist(_pd.to_numeric(self.df[num_cols[0]],errors='coerce').dropna()); ax.set_title('Histogram of '+num_cols[0]); self.canvas.draw()
                self.fig.savefig(png_path,dpi=150,bbox_inches='tight'); doc.add_heading('Figure',1); doc.add_picture(png_path,width=Inches(5))
            except Exception as e:
                doc.add_paragraph('(Plot not inserted: '+str(e)+')')
        doc.add_page_break(); doc.add_paragraph('StatMerge v1.0 — © 2025 Mirza Niaz Zaman Elin — MIT License'); doc.save(out_doc)
        QtWidgets.QMessageBox.information(self,'Report saved','Saved to '+out_doc)
